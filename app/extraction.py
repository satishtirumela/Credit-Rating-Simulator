"""
Server-side Gemini extraction engine for Credit Rating Simulator.
Extracts structured project input JSON from Key Input Template 1 (.docx) and Template 2 (.xlsx)
using prompt-only instruction grounding, and persists the extracted project data to Firestore.
"""

import json
import os
import time
from typing import Dict, Any, Optional
from dotenv import load_dotenv

load_dotenv()

SCHEMA_PATH = os.path.join(os.path.dirname(__file__), "..", "schemas", "project.schema.json")

EXTRACTION_SYSTEM_PROMPT = """You are a financial-statement extraction engine for a credit rating simulator.

You will be given two source documents for a single renewable-energy SPV:
1. A completed Key Input Template 1 (.docx) — Blocks A and D, notching inputs
2. A completed Key Input Template 2 (.xlsx) — Blocks B and C, DSCR schedule

Read both attached documents and return ONLY JSON matching the attached project.schema.json.

Rules:
- Return null for any field not present. NEVER return an empty string "" for a missing value — null and "" are not the same thing.
- Never infer or estimate a number or value. Never add fields not in the schema.
- Do not construct a name by combining other extracted fields (e.g. state + generic entity type). If the offtaker name field is blank or absent in the source, return null -- never synthesize a name from context.
- If dscr_schedule[] is populated with per-year data, you MUST return minimum_dscr and average_dscr as null. Do not calculate, estimate, or copy these values yourself under any circumstances when a schedule is present. They are derived downstream by the scoring engine, never by you.
- Do NOT include $schema, $id, or any other JSON Schema meta-fields in your output — only the actual field values themselves.
- Match enumerated fields on the exact stable CODE from Appendix A — never display prose.
- Extract offtakers[] individually (up to 4), each with name, type, contracted_share, rating_or_grade, rating_agency, rating_date, edition.
- rating_or_grade, rating_agency, and rating_date are three SEPARATE fields — never combine them into one string, even if the source document shows them together in a single cell or sentence. rating_or_grade must contain ONLY the bare rating symbol (e.g. "AA", "A+", "BB-") with nothing else appended. Example: if the source reads "AA, CRISIL, 12-03-2026", return rating_or_grade="AA", rating_agency="CRISIL", rating_date="2026-03-12" (reformatted to YYYY-MM-DD) — three fields, never rating_or_grade="AA, CRISIL, 12-03-2026".
- Itemise debt_instruments[] separately — never aggregate subordinated sponsor loans or CCDs into one line.
- p90_attestation is a single nested object with all four sub-fields (p90_plf, p90_attestation_basis, p90_resource_study, p90_preparer) — populate all four together or leave the object absent.
- Every percentage is a decimal fraction (0.9700 for 97%), never out of 100.
- Output JSON only. No commentary, no markdown fences."""


def _get_raw_schema() -> Dict[str, Any]:
    if os.path.exists(SCHEMA_PATH):
        with open(SCHEMA_PATH, "r", encoding="utf-8-sig") as sf:
            return json.load(sf)
    return {}


def save_to_firestore(project_id: str, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Saves extracted project JSON to Firestore under collection 'projects', document {project_id}.
    Falls back to local file storage if Firestore credentials are uninitialized.
    """
    from app.firestore import get_firestore_client_or_none

    safe_pid = project_id.strip() or "default_project"
    record = {
        "project_id": safe_pid,
        "extracted_data": extracted_data,
    }

    db = get_firestore_client_or_none()
    if db is not None:
        try:
            from firebase_admin import firestore
            doc_ref = db.collection("projects").document(safe_pid)
            doc_ref.set({
                "project_id": safe_pid,
                "extracted_data": extracted_data,
                "updated_at": firestore.SERVER_TIMESTAMP
            }, merge=True)
            return {"status": "saved", "provider": "firestore", "document_path": f"projects/{safe_pid}"}
        except Exception as err:
            local_cache_path = os.path.join(os.path.dirname(__file__), "..", "storage_uploads", f"firestore_{safe_pid}.json")
            os.makedirs(os.path.dirname(local_cache_path), exist_ok=True)
            with open(local_cache_path, "w", encoding="utf-8") as f:
                json.dump(record, f, indent=2)
            return {"status": "saved", "provider": f"local (firestore fallback: {str(err)})", "local_file": local_cache_path}

    local_cache_path = os.path.join(os.path.dirname(__file__), "..", "storage_uploads", f"firestore_{safe_pid}.json")
    os.makedirs(os.path.dirname(local_cache_path), exist_ok=True)
    with open(local_cache_path, "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2)
    return {"status": "saved", "provider": "local (firestore guard active)", "local_file": local_cache_path}


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extracts all paragraph and table text from a .docx file buffer."""
    import docx
    from io import BytesIO

    doc = docx.Document(BytesIO(docx_bytes))
    lines = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            lines.append(p.text.strip())

    for t_idx, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {t_idx + 1} ---")
        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(row_cells):
                lines.append(" | ".join(row_cells))

    return "\n".join(lines)


def _extract_xlsx_text(xlsx_bytes: bytes) -> str:
    """Extracts worksheet data and DSCR schedules from a .xlsx file buffer."""
    import openpyxl
    from io import BytesIO

    wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        lines.append(f"\n=== WORKSHEET: {sheet_name} ===")
        for row in sheet.iter_rows(values_only=True):
            if any(x is not None for x in row):
                row_str = " | ".join([str(x) if x is not None else "" for x in row])
                lines.append(row_str)

    return "\n".join(lines)


def extract_project_data(project_id: str, docx_bytes: bytes, xlsx_bytes: bytes, write_firestore: bool = True) -> Dict[str, Any]:
    """
    Extracts text from Template 1 (.docx) and Template 2 (.xlsx), calls Gemini API using prompt-only
    schema grounding at temperature=0.1, and optionally saves output to Firestore.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable is missing from server environment")

    from google import genai
    from google.genai import types

    client = genai.Client(api_key=api_key)
    raw_schema = _get_raw_schema()
    schema_str = json.dumps(raw_schema, indent=2)

    docx_text = _extract_docx_text(docx_bytes)
    xlsx_text = _extract_xlsx_text(xlsx_bytes)

    prompt = (
        "Here is the JSON Schema (project.schema.json) that your output MUST adhere to:\n\n"
        "```json\n"
        f"{schema_str}\n"
        "```\n\n"
        "Here are the two source documents for the renewable energy SPV:\n\n"
        "=== SOURCE DOCUMENT 1: Key Input Template 1 (.docx) ===\n"
        f"{docx_text}\n\n"
        "=== SOURCE DOCUMENT 2: Key Input Template 2 (.xlsx) ===\n"
        f"{xlsx_text}\n\n"
        "Read both attached source documents and return ONLY JSON matching project.schema.json."
    )

    config = types.GenerateContentConfig(
        system_instruction=EXTRACTION_SYSTEM_PROMPT,
        temperature=0.1,
        response_mime_type="application/json"
    )

    models_to_try = ["gemini-3.6-flash", "gemini-3.1-flash-lite", "gemini-2.5-flash"]
    response = None
    last_err = None

    for m_name in models_to_try:
        try:
            response = client.models.generate_content(
                model=m_name,
                contents=prompt,
                config=config
            )
            if response is not None:
                break
        except Exception as err:
            last_err = err
            err_str = str(err)
            print(f"[EXTRACTION WARNING] Gemini model {m_name} failed: {err_str}")
            # Only fall through to next model if error is 429 (rate limit), 503 (service unavailable), or 404 (model not found)
            if any(code in err_str for code in ["429", "503", "404", "RESOURCE_EXHAUSTED", "UNAVAILABLE", "NOT_FOUND"]):
                continue
            else:
                # Non-rate-limit error (e.g. malformed prompt/JSON or auth error) raises immediately
                raise RuntimeError(f"Gemini API model '{m_name}' error: {err_str}")

    if response is None:
        raise RuntimeError(f"Gemini API extraction failed across models {models_to_try}: {str(last_err)}")

    raw_text = response.text.strip()
    if raw_text.startswith("```json"):
        raw_text = raw_text[7:]
    if raw_text.startswith("```"):
        raw_text = raw_text[3:]
    if raw_text.endswith("```"):
        raw_text = raw_text[:-3]
    raw_text = raw_text.strip()

    extracted_json = json.loads(raw_text)

    # Post-processing safety net (CORE Appendix B note 1):
    # Whenever dscr_schedule[] is non-empty, minimum_dscr and average_dscr MUST be set to None/null
    # so they are derived solely by the Python scoring engine, never from model extraction.
    if isinstance(extracted_json, dict):
        if extracted_json.get("dscr_schedule"):
            extracted_json["minimum_dscr"] = None
            extracted_json["average_dscr"] = None

        # Offtaker name safety net: if offtaker name is copied from project_name, nullify it
        p_name = extracted_json.get("project_name")
        offtakers = extracted_json.get("offtakers")
        if isinstance(offtakers, list):
            for off in offtakers:
                if isinstance(off, dict):
                    off_name = off.get("name")
                    if off_name and (
                        (p_name and off_name.strip().lower() == p_name.strip().lower()) or
                        "discom" in off_name.lower() or
                        "spv" in off_name.lower()
                    ):
                        off["name"] = None

    firestore_res = None
    if write_firestore:
        firestore_res = save_to_firestore(project_id, extracted_json)

    return {
        "raw_text": raw_text,
        "extracted_data": extracted_json,
        "firestore_status": firestore_res
    }


def extract_project_from_storage(
    project_id: str,
    docx_filename: str,
    xlsx_filename: str,
    write_firestore: bool = True
) -> Dict[str, Any]:
    """
    Downloads template files directly FROM Firebase Storage bucket at projects/{project_id}/{filename},
    extracts structured project JSON using Gemini API, and saves output to Cloud Firestore.
    """
    from app.storage import get_project_file_bytes

    docx_bytes = get_project_file_bytes(project_id, docx_filename)
    xlsx_bytes = get_project_file_bytes(project_id, xlsx_filename)

    return extract_project_data(project_id, docx_bytes, xlsx_bytes, write_firestore=write_firestore)
